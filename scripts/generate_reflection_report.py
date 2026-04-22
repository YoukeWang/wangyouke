from __future__ import annotations

import html
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT_DIR / "reports"
DOCX_PATH = REPORT_DIR / "海外硕士录取标准分析模型_反思报告.docx"
FORMAT_CHECKLIST_PATH = REPORT_DIR / "格式符合项清单.md"
MANUAL_CHECKLIST_PATH = REPORT_DIR / "仍需人工确认项清单.md"

BLACK = RGBColor(0, 0, 0)
BODY_STYLE = "正文-学校规范"
CAPTION_STYLE = "图表题注-学校规范"
REFERENCE_STYLE = "参考文献-学校规范"


def clean_text(text: str) -> str:
    cleaned = text
    for old, new in {
        "$": "",
        "**": "",
        "`": "",
        "\\(": "(",
        "\\)": ")",
        "\\[": "[",
        "\\]": "]",
    }.items():
        cleaned = cleaned.replace(old, new)
    return cleaned


def set_element_fonts(element, east_asian: str, western: str) -> None:
    run_properties = element.get_or_add_rPr()
    fonts = run_properties.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.append(fonts)
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asian)
    fonts.set(qn("w:cs"), western)


def set_style_font(style, east_asian: str, western: str, size: float, bold: bool = False) -> None:
    style.font.name = western
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.underline = False
    style.font.color.rgb = BLACK
    set_element_fonts(style._element, east_asian, western)


def force_run_style(run, east_asian: str = "宋体", western: str = "Times New Roman") -> None:
    run.font.name = western
    run.font.color.rgb = BLACK
    run.font.underline = False
    set_element_fonts(run._element, east_asian, western)


def ensure_style(document: Document, name: str, style_type: WD_STYLE_TYPE):
    if name in document.styles:
        return document.styles[name]
    return document.styles.add_style(name, style_type)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    set_style_font(normal, "宋体", "Times New Roman", 12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    body = ensure_style(document, BODY_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = normal
    set_style_font(body, "宋体", "Times New Roman", 12)
    body.paragraph_format.first_line_indent = Pt(24)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_after = Pt(6)

    caption = ensure_style(document, CAPTION_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    caption.base_style = normal
    set_style_font(caption, "宋体", "Times New Roman", 10.5)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_after = Pt(6)

    reference = ensure_style(document, REFERENCE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    reference.base_style = normal
    set_style_font(reference, "宋体", "Times New Roman", 10.5)
    reference.paragraph_format.left_indent = Pt(21)
    reference.paragraph_format.first_line_indent = Pt(-21)
    reference.paragraph_format.line_spacing = 1.25
    reference.paragraph_format.space_after = Pt(3)

    title = document.styles["Title"]
    set_style_font(title, "黑体", "Times New Roman", 22, True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[style_name]
        set_style_font(style, "黑体", "Times New Roman", size, True)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    for style_name in ["TOC 1", "TOC 2", "TOC 3"]:
        if style_name in document.styles:
            toc_style = document.styles[style_name]
            set_style_font(toc_style, "宋体", "Times New Roman", 12)
            toc_style.paragraph_format.line_spacing = 1.25


def set_page_layout(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.gutter = Cm(0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


def set_page_numbering(section, start: int | None = None, fmt: str | None = None) -> None:
    page_number_type = section._sectPr.find(qn("w:pgNumType"))
    if page_number_type is None:
        page_number_type = OxmlElement("w:pgNumType")
        section._sectPr.append(page_number_type)
    if start is not None:
        page_number_type.set(qn("w:start"), str(start))
    if fmt is not None:
        page_number_type.set(qn("w:fmt"), fmt)


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_complex_field(paragraph, instruction: str, placeholder: str = "") -> None:
    begin_run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(field_begin)

    instruction_run = paragraph.add_run()
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    instruction_run._r.append(instruction_text)

    separate_run = paragraph.add_run()
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(field_separate)

    if placeholder:
        result_run = paragraph.add_run(placeholder)
        force_run_style(result_run)

    end_run = paragraph.add_run()
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(field_end)


def set_header_footer(section, header_text: str = "", numbered: bool = False) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header_paragraph = section.header.paragraphs[0]
    clear_paragraph(header_paragraph)
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header_text:
        run = header_paragraph.add_run(clean_text(header_text))
        run.font.size = Pt(10.5)
        force_run_style(run)

    footer_paragraph = section.footer.paragraphs[0]
    clear_paragraph(footer_paragraph)
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if numbered:
        add_complex_field(footer_paragraph, "PAGE")
        for run in footer_paragraph.runs:
            run.font.size = Pt(10.5)
            force_run_style(run)


def set_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_paragraph(document: Document, text: str, style: str = BODY_STYLE, alignment=None):
    paragraph = document.add_paragraph(clean_text(text), style=style)
    if alignment is not None:
        paragraph.alignment = alignment
    for run in paragraph.runs:
        force_run_style(run)
    return paragraph


def add_heading(document: Document, text: str, level: int = 1, centered: bool = False):
    paragraph = document.add_paragraph(clean_text(text), style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        force_run_style(run, "黑体", "Times New Roman")
        run.font.bold = True
    return paragraph


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(clean_text(text), style=CAPTION_STYLE)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        force_run_style(run)
        run.font.size = Pt(10.5)


def set_cell_text(cell, text: str, bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = alignment
    run = paragraph.add_run(clean_text(text))
    run.font.bold = bold
    run.font.size = Pt(10.5)
    force_run_style(run)


def add_field_table(document: Document) -> None:
    add_caption(document, "表 1 样例数据主要字段说明")
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["字段", "含义", "作用"]
    rows = [
        ["min_gpa_100", "最低 GPA 门槛", "判断申请人是否满足学术硬门槛"],
        ["min_ielts / min_toefl", "语言要求", "判断雅思或托福是否达标"],
        ["gre_policy / min_gre", "GRE 要求", "区分不要求、可选和必需"],
        ["major_requirement", "专业背景要求", "评估本科专业匹配度"],
        ["quant_requirement", "数学/统计要求", "评估量化基础是否足够"],
        ["programming_requirement", "编程要求", "评估技术基础是否足够"],
        ["internship_importance", "实习重要度", "评估职业导向项目的软实力"],
        ["research_importance", "科研重要度", "评估研究导向项目的软实力"],
        ["tuition_usd", "项目总学费", "判断是否符合预算"],
    ]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, alignment=WD_ALIGN_PARAGRAPH.LEFT)


def remove_table_borders(table) -> None:
    table_properties = table._tbl.tblPr
    borders = table_properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "nil")


def add_formula(document: Document) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    set_cell_text(
        table.cell(0, 0),
        "录取匹配指数 = 50%×硬门槛 + 30%×专业匹配 + 20%×软实力 + 预算修正 − 竞争度惩罚",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_text(table.cell(0, 1), "(1)", alignment=WD_ALIGN_PARAGRAPH.RIGHT)


def add_cover(document: Document) -> None:
    set_page_layout(document.sections[0])
    set_header_footer(document.sections[0], numbered=False)
    document.add_paragraph()
    document.add_paragraph()
    title = document.add_paragraph("海外硕士录取标准分析模型", style="Title")
    subtitle = document.add_paragraph("反思报告", style="Title")
    for paragraph in [title, subtitle]:
        for run in paragraph.runs:
            force_run_style(run, "黑体", "Times New Roman")
            run.font.size = Pt(22)
            run.font.bold = True

    document.add_paragraph()
    cover_lines = [
        "赛道：交互式数据分析工具",
        "项目名称：海外硕士录取标准分析模型",
        "提交内容：反思报告 Word 终稿",
        "姓名：请填写",
        "学号：请填写",
        "班级：请填写",
        "日期：2026 年 4 月",
    ]
    for line in cover_lines:
        paragraph = document.add_paragraph(clean_text(line), style=BODY_STYLE)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            force_run_style(run)


def add_front_matter(document: Document) -> None:
    front_section = document.add_section(WD_SECTION.NEW_PAGE)
    set_page_layout(front_section)
    set_page_numbering(front_section, start=1, fmt="lowerRoman")
    set_header_footer(front_section, "海外硕士录取标准分析模型反思报告", numbered=True)

    add_heading(document, "摘要", level=1, centered=True)
    add_paragraph(
        document,
        "本反思报告围绕海外硕士录取标准分析模型的设计与实现展开。项目选择交互式数据分析工具赛道，使用 Python、Streamlit、pandas 和 Altair 构建面向留学申请场景的小型数据产品。工具将录取标准拆分为硬门槛、专业匹配和软实力三个维度，通过样例项目数据展示不同国家和专业方向的申请要求，并根据申请人画像输出保底、匹配、冲刺和暂不达标四类结果。",
    )
    keywords = document.add_paragraph(
        "关键词：留学申请；录取标准；交互式数据分析；Streamlit", style=BODY_STYLE
    )
    keywords.paragraph_format.first_line_indent = Pt(0)
    for run in keywords.runs:
        force_run_style(run)

    document.add_page_break()
    add_heading(document, "ABSTRACT", level=1, centered=True)
    add_paragraph(
        document,
        "This reflection report reviews the design and implementation of an overseas master's admission analysis model. The project follows the interactive data analysis tool track and uses Python, Streamlit, pandas, and Altair to build a lightweight data product for study abroad planning. The model divides admission standards into hard requirements, academic and major fit, and soft power, then provides interpretable recommendation levels based on a sample applicant profile.",
    )
    keywords_en = document.add_paragraph(
        "Keywords: study abroad; admission standards; interactive data analysis; Streamlit",
        style=BODY_STYLE,
    )
    keywords_en.paragraph_format.first_line_indent = Pt(0)
    for run in keywords_en.runs:
        force_run_style(run)

    document.add_page_break()
    add_heading(document, "目录", level=1, centered=True)
    toc = document.add_paragraph()
    toc.paragraph_format.first_line_indent = Pt(0)
    add_complex_field(toc, r'TOC \o "1-3" \h \z \u', "目录将在 Word 中自动更新。")
    for run in toc.runs:
        force_run_style(run)


def add_body(document: Document) -> None:
    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    set_page_layout(body_section)
    set_page_numbering(body_section, start=1, fmt="decimal")
    set_header_footer(body_section, "海外硕士录取标准分析模型反思报告", numbered=True)

    add_heading(document, "1 项目概述", level=1)
    add_heading(document, "1.1 选题背景", level=2)
    add_paragraph(
        document,
        "我选择海外硕士录取标准分析作为作业主题，是因为留学申请本身就是一个多指标决策问题。申请者不仅需要关注学校排名，还要同时比较 GPA、语言成绩、GRE、专业背景、实习科研经历和费用预算。传统表格可以保存信息，但不方便根据个人背景实时筛选和解释，因此适合做成交互式数据分析工具。",
    )
    add_paragraph(
        document,
        "项目重点放在商科和数据类硕士项目上，包括 Business Analytics、Data Science、Information Systems、Statistics 等方向。这些方向通常既关注学术门槛，也关注专业基础和实践经历，能够较好体现硬门槛、专业匹配和软实力三类因素。",
    )

    add_heading(document, "1.2 产品目标", level=2)
    add_paragraph(
        document,
        "本工具的目标不是替代真实申请咨询，而是帮助用户完成初步选校分层。用户输入本科背景、GPA、语言成绩、GRE、专业匹配度、数学和编程基础、实习科研经历以及预算后，系统会根据样例项目数据计算录取匹配指数，并输出保底、匹配、冲刺和暂不达标四类结果。",
    )

    add_heading(document, "2 数据与模型设计", level=1)
    add_heading(document, "2.1 数据结构", level=2)
    add_paragraph(
        document,
        "项目使用小规模样例数据，共整理二十二个海外硕士项目，覆盖英国、美国、新加坡、中国香港、澳大利亚、加拿大和荷兰。数据字段包括学校、国家或地区、项目名称、专业方向、排名、学费、学制、GPA 门槛、语言要求、GRE 要求、专业背景要求、数学统计要求、编程要求、实习重要度和科研重要度。",
    )
    add_field_table(document)

    add_heading(document, "2.2 模型逻辑", level=2)
    add_paragraph(
        document,
        "模型采用规则加权方式，而不是复杂机器学习模型。这样做的原因是样例数据量较小，规则模型更透明，也更适合课堂演示。模型首先检查硬门槛是否满足，然后计算三个核心维度得分，并结合预算和竞争度进行修正。",
    )
    add_formula(document)
    add_paragraph(
        document,
        "硬门槛主要包括 GPA、语言成绩和 GRE。专业匹配主要包括本科专业相关度、数学统计基础和编程基础。软实力主要包括实习时长、科研经历和数据分析项目经历。最终分数越高，说明申请人与项目要求越匹配。",
    )

    add_heading(document, "2.3 交互设计", level=2)
    add_paragraph(
        document,
        "页面使用 Streamlit 搭建。侧边栏用于输入申请人画像，主页面展示项目门槛、匹配结果、模型解释、费用对比和样例数据。用户调整任意条件后，页面会自动重新计算结果，这体现了交互式数据产品相较静态报告的优势。",
    )

    add_heading(document, "3 分析结果与应用价值", level=1)
    add_heading(document, "3.1 分析结果", level=2)
    add_paragraph(
        document,
        "通过项目门槛图表可以看到，不同国家和项目方向之间的 GPA、语言和费用要求存在明显差异。排名靠前或竞争更强的项目通常对 GPA、语言成绩和专业背景提出更高要求。对于数据科学、统计和应用计算方向，数学统计和编程基础的重要性也更加突出。",
    )
    add_paragraph(
        document,
        "个人匹配页面能够把申请结果分为保底、匹配、冲刺和暂不达标。相比简单罗列学校信息，这种分层结果更接近真实选校场景，也便于在演示视频中解释为什么同一个申请人面对不同项目会得到不同建议。",
    )

    add_heading(document, "3.2 产品价值", level=2)
    add_paragraph(
        document,
        "本工具的价值主要体现在三个方面。第一，它把复杂的申请标准转化为可视化和可交互的界面。第二，它通过模型解释页面展示每个项目的得分拆解，使结果不是黑箱。第三，它能够帮助用户快速识别短板，例如语言未达标、专业背景不足、编程基础较弱或预算压力较大。",
    )

    add_heading(document, "4 反思与改进", level=1)
    add_heading(document, "4.1 完成过程的收获", level=2)
    add_paragraph(
        document,
        "完成这个项目后，我更加理解数据产品和普通数据分析报告的区别。数据报告通常是一次性输出结论，而交互式工具需要考虑用户输入、结果更新、可解释性和页面展示逻辑。为了让工具更像产品，我没有只展示表格，而是设计了申请人画像、匹配指数、结果分层和单项目解释。",
    )
    add_paragraph(
        document,
        "我也认识到，在样例数据较少的情况下，透明的规则模型比复杂模型更合适。虽然规则模型不能给出真实录取概率，但它能够清楚说明每个判断来自哪些维度，更适合作为课程作业和产品原型展示。",
    )

    add_heading(document, "4.2 存在不足", level=2)
    add_paragraph(
        document,
        "项目目前仍然存在不足。首先，数据为课堂演示样例，并非实时爬取的官方数据，因此不能直接用于真实申请决策。其次，模型权重由人工设定，虽然有解释性，但不一定完全符合不同学校的真实审核逻辑。再次，项目没有考虑文书质量、推荐信、面试表现和申请时间等重要因素。",
    )

    add_heading(document, "4.3 后续改进方向", level=2)
    add_paragraph(
        document,
        "后续可以从三个方向继续改进。第一，扩充数据来源，加入更多学校和项目，并标注数据更新时间。第二，增加可配置权重功能，让用户根据商科、数据科学或信息系统等不同方向调整模型。第三，如果能够收集历史申请案例，可以进一步训练简单的分类模型，用真实案例校正规则模型。",
    )

    add_heading(document, "5 总结", level=1)
    add_paragraph(
        document,
        "总体来看，本项目完成了赛道四对交互式数据分析工具的要求：有可运行的 Python 应用，有数据清洗和指标计算，有交互筛选、可视化图表和可解释的模型结果。项目规模不大，但主题明确，能够围绕留学申请场景展示数据产品从数据、模型到页面交互的完整流程。",
    )


def add_back_matter(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "参考文献", level=1, centered=True)
    references = [
        "[1] Streamlit Team. Streamlit Documentation.",
        "[2] The pandas development team. pandas Documentation.",
        "[3] Altair Developers. Altair Documentation.",
        "[4] Python Software Foundation. Python Documentation.",
        "[5] 课程作业任务说明：交互式数据分析工具赛道要求.",
    ]
    for item in references:
        paragraph = document.add_paragraph(clean_text(item), style=REFERENCE_STYLE)
        for run in paragraph.runs:
            force_run_style(run)

    document.add_page_break()
    add_heading(document, "附录 A 项目运行说明", level=1, centered=True)
    add_paragraph(document, "安装依赖命令：pip install -e .")
    add_paragraph(document, "启动应用命令：streamlit run main.py")
    add_paragraph(document, "本地访问地址：http://localhost:8501")

    document.add_page_break()
    add_heading(document, "致谢", level=1, centered=True)
    add_paragraph(
        document,
        "感谢课程提供交互式数据分析工具赛道，使我能够以产品化方式完成 Python 数据分析实践。也感谢项目开发过程中使用到的开源工具，它们帮助我更快完成从数据整理、模型设计到可视化展示的完整流程。",
    )


def build_document() -> Document:
    document = Document()
    configure_styles(document)
    set_update_fields_on_open(document)
    add_cover(document)
    add_front_matter(document)
    add_body(document)
    add_back_matter(document)
    return document


def update_fields_with_word(docx_path: Path) -> tuple[bool, str]:
    try:
        import win32com.client  # type: ignore
    except Exception as error:
        return False, f"未启用 Word COM 自动刷新：{error}"

    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(docx_path.resolve()))
        document.Fields.Update()
        for table_of_contents in document.TablesOfContents:
            table_of_contents.Update()
        document.Repaginate()
        document.Save()
        document.Close(False)
        word.Quit()
        return True, "已通过 Word COM 自动刷新目录、页码域和交叉字段。"
    except Exception as error:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        return False, f"Word COM 自动刷新失败：{error}"


def extract_docx_text(docx_path: Path) -> str:
    with ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    text_parts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", document_xml)
    return html.unescape("".join(text_parts))


def inspect_docx_xml(docx_path: Path) -> dict:
    with ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
    full_text = extract_docx_text(docx_path)
    residual_symbols = [symbol for symbol in ["$", "**", "`"] if symbol in full_text]
    return {
        "has_toc_field": "TOC" in document_xml,
        "has_update_fields": "updateFields" in settings_xml,
        "has_lower_roman": 'w:fmt="lowerRoman"' in document_xml,
        "has_decimal": 'w:fmt="decimal"' in document_xml,
        "has_no_markdown_residue": not residual_symbols,
        "residual_symbols": residual_symbols,
    }


def write_checklists(com_updated: bool, com_message: str, inspection: dict) -> None:
    rows = [
        ("页面设置", "A4；上、下 2.5 厘米，左 3.0 厘米，右 2.5 厘米；装订线 0 厘米；页眉 1.5 厘米，页脚 1.75 厘米。", "满足"),
        ("分节页码", "封面不编号；摘要、英文摘要、目录使用小写罗马数字；正文起使用阿拉伯数字并从 1 开始。", "满足"),
        ("自动目录", "已插入 Word 自动目录域，并设置打开文件时自动更新字段。", "满足" if inspection["has_toc_field"] and inspection["has_update_fields"] else "需复查"),
        ("目录刷新", com_message, "满足" if com_updated else "需人工确认"),
        ("标题体系", "正文标题采用 1、1.1、1.1.1 兼容体系。", "满足"),
        ("标题颜色", "Title、Heading 1、Heading 2、Heading 3 和 TOC 样式均强制为黑色。", "满足"),
        ("正文格式", "正文使用宋体小四，1.5 倍行距，首行缩进 2 字符。", "满足"),
        ("摘要格式", "包含中文摘要、中文关键词、英文摘要、英文关键词，且独立分页。", "满足"),
        ("关键词数量", "中文关键词 4 个，英文关键词 4 个，符合 3 至 5 个的一般要求。", "满足"),
        ("Markdown 残留", f"检测残留符号：{inspection['residual_symbols']}", "满足" if inspection["has_no_markdown_residue"] else "不满足"),
        ("表格格式", "表题在表上，表格居中，内容按学校论文常用格式设置。", "满足"),
        ("公式格式", "公式居中，编号右对齐。", "满足"),
        ("参考文献", "参考文献按 [1] 至 [5] 顺序编号。", "满足"),
        ("附录与致谢", "附录 A 和致谢均独立分页。", "满足"),
    ]
    lines = [
        "# 格式符合项清单",
        "",
        "说明：用户未提供学校格式要求文件、封面模板、声明页模板和图片目录，因此本报告按通用学校论文 Word 终稿规范生成。",
        "",
        "| 项目 | 处理结果 | 状态 |",
        "| --- | --- | --- |",
    ]
    lines.extend(f"| {name} | {result} | {status} |" for name, result, status in rows)
    FORMAT_CHECKLIST_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")

    manual_lines = [
        "# 仍需人工确认项清单",
        "",
        "- 学校未提供具体格式文件：请确认学院是否另有页边距、字体、封面或参考文献细则。",
        "- 封面模板未提供：当前使用通用封面，请填写姓名、学号、班级。",
        "- 声明页模板未提供：当前未插入声明页，如学校要求需另行加入。",
        "- 图片目录未提供：当前报告没有插入应用截图，如教师要求可补充截图和图题。",
        "- 样例数据不是官方实时数据：真实申请前必须以学校官网为准。",
        "- 演示视频由你自行录制：建议展示申请画像输入、匹配结果、模型解释和费用对比。",
    ]
    if not com_updated:
        manual_lines.append("- 本机未完成 Word COM 自动刷新：首次打开 Word 后请按 Ctrl+A，再按 F9 更新目录和页码域。")
    MANUAL_CHECKLIST_PATH.write_text("\n".join(manual_lines) + "\n", encoding="utf-8")


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(DOCX_PATH)
    com_updated, com_message = update_fields_with_word(DOCX_PATH)
    inspection = inspect_docx_xml(DOCX_PATH)
    write_checklists(com_updated, com_message, inspection)
    print(f"Word 文件：{DOCX_PATH}")
    print(f"格式清单：{FORMAT_CHECKLIST_PATH}")
    print(f"人工确认：{MANUAL_CHECKLIST_PATH}")
    print(com_message)
    if inspection["residual_symbols"]:
        print(f"发现残留符号：{inspection['residual_symbols']}")
    else:
        print("未发现 Markdown 或 LaTeX 定界符残留。")


if __name__ == "__main__":
    main()
